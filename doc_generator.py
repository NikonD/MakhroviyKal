"""
Генерация итогового docx «Экспертное заключение».
Структура таблицы повторяет колонки из примера (без сложного форматирования):

  №, ФИО, ОП, Курс,
  По плану: Дисциплина, Кредиты,
  Фактически: Курсы (название + источник + часы), Кредиты/часы, Баллы,
  Заключение: Соответствие, Примечание,
  Зачтённая оценка
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


# ---------------------------------------------------------------------------
# Public payload schema
# ---------------------------------------------------------------------------
@dataclass
class CourseEntry:
    course_title: str = ""
    provider: str = ""
    hours: str = ""
    grade: str = ""
    date: str = ""

    @property
    def clean_hours(self) -> str:
        if not self.hours.strip():
            return ""
        m = re.search(r"(\d+(?:[.,]\d+)?)", self.hours)
        if m:
            try:
                val = float(m.group(1).replace(",", "."))
                return f"{int(val)} ч" if val.is_integer() else f"{val} ч"
            except ValueError:
                pass
        return self.hours.strip()

    def format(self) -> str:
        """«Название курса (Провайдер, часы)» — вместо даты выводим часы."""
        head = self.course_title.strip()
        meta_bits = [b for b in (self.provider.strip(), self.clean_hours) if b]
        if meta_bits:
            head = f"{head} ({', '.join(meta_bits)})" if head else f"({', '.join(meta_bits)})"
        return head


@dataclass
class DisciplineRow:
    discipline: str = ""
    plan_credits: str = ""
    courses: list[CourseEntry] = field(default_factory=list)
    total_hours: str = ""
    grade_points: str = ""
    compliance: str = "полное"
    note: str = ""
    final_grade: str = ""


@dataclass
class ReportPayload:
    student_name: str = ""
    program_code: str = ""
    course_year: str = ""
    rows: list[DisciplineRow] = field(default_factory=list)
    experts: list[str] = field(default_factory=lambda: [
        "Борамбаева Г.М., ст.преподаватель кафедры ИКТ",
        "Жекеева С.С., ст.преподаватель кафедры ИКТ",
        "Никишина О.А., ст.преподаватель кафедры ИКТ",
    ])


# ---------------------------------------------------------------------------
# Helpers: derive program name from group code (e.g. АПО-23 -> Архитектор ПО)
# ---------------------------------------------------------------------------
PROGRAM_MAP: dict[str, str] = {
    "АПО": "6B06105 «Архитектор программного обеспечения»",
    "ИС":  "6B06104 «Информационные системы»",
    "ВТ":  "6B06103 «Вычислительная техника и программное обеспечение»",
    "ИВТ": "6B06103 «Вычислительная техника и программное обеспечение»",
    "ИКТ": "6B06102 «Информационно-коммуникационные технологии»",
}


def derive_program(group: str) -> str:
    if not group:
        return ""
    m = re.match(r"\s*([А-Яа-яA-Za-z]+)", group)
    if not m:
        return ""
    prefix = m.group(1).upper()
    return PROGRAM_MAP.get(prefix, "")


def short_program_name(code_and_title: str) -> str:
    """`6B06104 «Информационные системы»` -> `Информационные системы`."""
    if not code_and_title:
        return ""
    m = re.search(r"[«\"]([^«»\"]+)[»\"]", code_and_title)
    if m:
        return m.group(1).strip()
    return code_and_title.strip()


def format_student_name(name: str) -> str:
    """Форматирует ФИО как в референсе: «Фамилия Имя\\nОтчество».

    Понимает оба порядка: «Фамилия Имя Отчество» (типично для рус. документов)
    и «Имя Фамилия» (часто возвращает OCR).
    Сокращения вроде «Сычёв В.А.» оставляем как есть.
    """
    if not name:
        return ""
    s = re.sub(r"\s+", " ", name).strip()
    if "." in s:
        return s
    parts = s.split(" ")
    if len(parts) == 3:
        return f"{parts[0]} {parts[1]}\n{parts[2]}"
    if len(parts) == 2:
        # Эвристика: славянские фамилии часто оканчиваются на -ов/-ев/-ин/-ский/-ёв/-ук
        surn_suffixes = ("ов", "ев", "ёв", "ин", "ын", "ий", "ский", "цкий", "ук", "юк", "енко", "ко")
        if parts[1].lower().endswith(surn_suffixes) and not parts[0].lower().endswith(surn_suffixes):
            return f"{parts[1]} {parts[0]}"
    return s


# ---------------------------------------------------------------------------
# Row helpers: hours / grades aggregation
# ---------------------------------------------------------------------------
def sum_hours_text(entries: list[CourseEntry]) -> str:
    """Сумма часов первых чисел из hours каждого курса. Для строки с одной дисциплиной."""
    total = 0
    found = False
    for e in entries:
        m = re.search(r"(\d+(?:[.,]\d+)?)", e.hours)
        if m:
            try:
                total += int(float(m.group(1).replace(",", ".")))
                found = True
            except ValueError:
                pass
    return f"{total} ч" if found else ""


def list_hours_text(entries: list[CourseEntry]) -> str:
    """Часы каждого курса по строкам. Используется в строке «не сопоставлено»."""
    lines = [e.hours.strip() or "—" for e in entries if e.course_title.strip()]
    return "\n".join(lines)


# def list_grades_text(entries: list[CourseEntry]) -> str:
#     """Баллы/оценки каждого курса по строкам."""
#     lines = [e.grade.strip() or "—" for e in entries if e.course_title.strip()]
#     return "\n".join(lines)


# ---------------------------------------------------------------------------
# DOCX rendering
# ---------------------------------------------------------------------------
# Заголовки пишем ТОЛЬКО в «ведущие» ячейки (после мерджа).
# Иначе при merge() Word склеивает текст из всех донорских ячеек и шапка дублируется.
# Ключ — индекс «ведущей» (top-left) ячейки.
HEADERS_ROW_1_AT: dict[int, str] = {
    0: "№",
    1: "ФИО обучающегося",
    2: "Код и наименование ОП",
    3: "Курс",
    4: "По учебному плану СКУ им.М.Козыбаева",       # объединена 4-5
    6: "Фактически изучено претендентом и рассмотрено экспертной группой",  # 6-8
    9: "Заключение экспертной группы",               # 9-10
    11: "Зачтенная оценка по дисциплине учебного плана",
}
HEADERS_ROW_2_AT: dict[int, str] = {
    4: "Наименование перезачитываемых дисциплин и видов учебной работы",
    5: "Количество кредитов",
    6: "Наименование курса неформального обучения (наименование, номер, количество часов)",
    7: "Количество кредитов (часов)",
    8: "Баллы (оценка)",
    9: "Соответствие целей программы обучения, объёма программы и оценки (полное, частичное, не соответствует)",
    10: "Примечание",
}
MERGE_SPANS = [
    # (row_idx, start_col, end_col) — групповые шапки в row0
    (0, 4, 5),    # По учебному плану
    (0, 6, 8),    # Фактически
    (0, 9, 10),   # Заключение
]
MERGE_DOWN_COLS = [0, 1, 2, 3, 11]  # колонки, где row0 объединяется с row1


# Шрифт во всём документе — Times New Roman 12pt
FONT_NAME = "Times New Roman"
TABLE_FONT_SIZE_PT = 12
HEADER_FONT_SIZE_PT = 12

# Ширины колонок (см) — пропорции из ЭЗ_Умербеков_Н_ИС-у-25.docx,
# отмасштабированные под A4 landscape с полями 2/2 (доступно 25.7 см)
# Ширины 1-в-1 как в референсе ЭЗ_Умербеков_Н_ИС-у-25.docx (сумма 28.25 см).
# Чтобы таблица влезла без обрезания, поля страницы делаем 1.0 см.
COL_WIDTHS_CM: list[float] = [
    0.75,  # 0: №
    3.38,  # 1: ФИО
    2.75,  # 2: ОП
    1.00,  # 3: Курс
    4.00,  # 4: Дисциплина
    1.25,  # 5: Кредиты
    5.31,  # 6: Курсы
    1.19,  # 7: Часы
    2.00,  # 8: Баллы
    3.25,  # 9: Соответствие
    1.00,  # 10: Примечание
    2.37,  # 11: Зачтённая оценка
]

# Колонки данных, которые выравниваем по центру (короткие числовые поля).
# Остальные — по левому краю.
DATA_ALIGN_CENTER = {0, 3, 5, 7, 8, 9, 10, 11}

def _apply_font(run, *, bold: bool, size_pt: int) -> None:
    """Times New Roman, нужный размер, опц. жирный — включая cs/eastAsia, чтобы Word везде применил."""
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.bold = bold
    # python-docx по умолчанию ставит только w:rFonts/@w:ascii. Чтобы шрифт
    # точно применился к кириллице, дублируем cs/hAnsi/eastAsia.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    align_center: bool = True,  # по умолчанию все ячейки таблицы — по центру
    size_pt: int = TABLE_FONT_SIZE_PT,
) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Все строки идут в ОДНОМ параграфе через line-break — иначе Word
    # добавляет вертикальные отступы между параграфами и ячейка раздувается.
    lines = str(text).split("\n")
    for li, line in enumerate(lines):
        if li > 0:
            br_run = para.add_run()
            _apply_font(br_run, bold=bold, size_pt=size_pt)
            br_run.add_break()
        run = para.add_run(line)
        _apply_font(run, bold=bold, size_pt=size_pt)


def _set_landscape(section) -> None:
    # Явно A4 landscape (29.7 × 21 см), а не US Letter
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    # Поля 1 см / 1 см L/R чтобы таблица 28.25 см влезла без обрезания.
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)


def _set_col_widths(table, widths_cm: list[float]) -> None:
    """Жёстко задаём ширину каждой колонки во всех строках (после merges)."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])
    # Дополнительно: фиксируем layout=fixed чтобы Word не перерастягивал
    from docx.oxml import OxmlElement
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _add_borders(table) -> None:
    from docx.oxml import OxmlElement
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:color"), "000000")
            borders.append(el)
        tbl_pr.append(borders)


def build_report(payload: ReportPayload) -> bytes:
    doc = Document()
    _set_landscape(doc.sections[0])

    # Дефолтный шрифт документа — Times New Roman 12pt
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(TABLE_FONT_SIZE_PT)
    # Для надёжности дублируем имя шрифта в rPrDefault
    from docx.oxml import OxmlElement
    rpr_default = style.element.get_or_add_rPr()
    rfonts = rpr_default.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr_default.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)

    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Экспертное заключение")
    _apply_font(run, bold=True, size_pt=12)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        "по определению соответствия результатов неформального обучения по профилю подготовки"
    )
    _apply_font(sub_run, bold=True, size_pt=12)

    doc.add_paragraph("")

    rows_data = payload.rows or [DisciplineRow()]

    table = doc.add_table(rows=2 + len(rows_data), cols=12)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # СНАЧАЛА мерджим (чтобы не дублировался текст), ПОТОМ пишем
    for r, s, e in MERGE_SPANS:
        table.rows[r].cells[s].merge(table.rows[r].cells[e])
    for col in MERGE_DOWN_COLS:
        table.rows[0].cells[col].merge(table.rows[1].cells[col])

    # Заголовки — только в ведущие ячейки
    for ci, h in HEADERS_ROW_1_AT.items():
        _set_cell_text(
            table.rows[0].cells[ci], h, align_center=True,
            size_pt=HEADER_FONT_SIZE_PT,
        )
    for ci, h in HEADERS_ROW_2_AT.items():
        _set_cell_text(
            table.rows[1].cells[ci], h, align_center=True,
            size_pt=HEADER_FONT_SIZE_PT,
        )

    student_name_fmt = format_student_name(payload.student_name)
    program_short = short_program_name(payload.program_code)

    for ri, row_data in enumerate(rows_data, start=2):
        cells = table.rows[ri].cells
        # Группой считаем строку с >1 курсом (типичный случай — «не сопоставлено»)
        is_group = len(row_data.courses) > 1
        courses_text = "\n".join(
            c.format() for c in row_data.courses if c.course_title.strip()
        ) or ""
        compliance = row_data.compliance or ("" if is_group else "полное")
        if is_group:
            hours_cell = row_data.total_hours or sum_hours_text(row_data.courses)
            #grade_cell = row_data.grade_points or list_grades_text(row_data.courses)
            grade_cell=""
        else:
            hours_cell = row_data.total_hours or sum_hours_text(row_data.courses)
            grade_cell = row_data.grade_points or (
            row_data.courses[0].grade.strip() if row_data.courses else ""
            )

        values = [
            str(ri - 1),           # 0  №
            student_name_fmt,      # 1  ФИО
            program_short,         # 2  ОП
            payload.course_year,   # 3  Курс
            row_data.discipline,   # 4  Дисциплина
            row_data.plan_credits, # 5  Кредиты
            courses_text,          # 6  Курсы
            hours_cell,            # 7  Часы
            grade_cell,            # 8  Баллы
            compliance,            # 9  Соответствие
            row_data.note,         # 10 Примечание
            row_data.final_grade,  # 11 Зачт. оценка
        ]
        for ci, v in enumerate(values):
            _set_cell_text(cells[ci], v, align_center=(ci in DATA_ALIGN_CENTER))

    _set_col_widths(table, COL_WIDTHS_CM)
    _add_borders(table)

    # Подписи экспертов — формат как в референсе: длинные с табами/nbsp
    doc.add_paragraph("")
    expert_head = doc.add_paragraph()
    expert_head.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _apply_font(expert_head.add_run("Эксперты:"), bold=True, size_pt=TABLE_FONT_SIZE_PT)
    nbsp = "\u00a0"
    for expert in payload.experts:
        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = line.add_run(
            f"{expert}\t{nbsp*7}______________{nbsp*4}______________"
        )
        _apply_font(run, bold=False, size_pt=TABLE_FONT_SIZE_PT)
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        sub_run = sub.add_run(
            f"\t\t\t(ФИО, должность)\t\t\t\t\t{nbsp*10}(подпись)\t{nbsp*27}(дата)"
        )
        _apply_font(sub_run, bold=False, size_pt=10)
        sub_run.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
